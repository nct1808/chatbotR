def __import os
import json
import pickle
import logging
from typing import List, Dict, Any, Optional
from datetime import datetime
import streamlit as st
from dotenv import load_dotenv
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
import openai
from openai import OpenAI
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings
from langchain.docstore.document import Document
import io
from docx import Document as DocxDocument
import PyPDF2
import pandas as pd
import hashlib
import time
import re
import numpy as np

# Load environment variables
load_dotenv()

# API KEY - Sử dụng biến môi trường (An toàn)
OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')

if not OPENAI_API_KEY:
    raise ValueError("❌ OPENAI_API_KEY not found. Please set it in .env file")

# Cấu hình logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Constants
SCOPES = ['https://www.googleapis.com/auth/drive.readonly']
CREDENTIALS_FILE = 'credentials.json'
TOKEN_FILE = 'token.pickle'
VECTOR_STORE_DIR = 'vector_store'
CACHE_DIR = 'cache'
EMBEDDING_CACHE_DIR = 'embedding_cache'
TRANSLATION_CACHE_DIR = 'translation_cache'

class GoogleDriveRAGChatbot:
    def __init__(self, openai_api_key: Optional[str] = None):
        """Khởi tạo chatbot RAG với OpenAI API"""
        # Use environment variable API key
        self.openai_api_key = openai_api_key or OPENAI_API_KEY
        
        if not self.openai_api_key:
            raise ValueError("OpenAI API Key not configured")
        
        # Cấu hình OpenAI
        self.client = OpenAI(api_key=self.openai_api_key)
        self.embeddings = OpenAIEmbeddings(
            model="text-embedding-3-large",
            openai_api_key=self.openai_api_key
        )
        
        # Cấu hình
        self.chunk_size = int(os.getenv('CHUNK_SIZE', 500))
        self.chunk_overlap = int(os.getenv('CHUNK_OVERLAP', 100))
        self.max_search_results = int(os.getenv('MAX_SEARCH_RESULTS', 5))
        self.max_file_size_mb = int(os.getenv('MAX_FILE_SIZE_MB', 50))
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len
        )
        
        # Separate vector stores for different sources
        self.drive_vector_store = None
        self.upload_vector_store = None
        self.current_source = "drive"  # "drive" or "upload"
        
        # Drive service
        self.drive_service = None
        
        # Cache systems
        self.file_cache = {}
        self.embedding_cache = {}
        self.translation_cache = {}
        self.vector_metadata = {}
        
        # Memory and Learning System
        self.conversation_memory = []
        self.user_preferences = {
            'preferred_answer_length': 'medium',  # short, medium, long
            'preferred_language': 'vi',
            'frequently_asked_topics': {},
            'user_interaction_patterns': []
        }
        self.learning_data = {
            'successful_queries': [],
            'failed_queries': [],
            'user_feedback': {},
            'context_patterns': {}
        }
        
        # Tạo cache directories
        for cache_dir in [CACHE_DIR, EMBEDDING_CACHE_DIR, TRANSLATION_CACHE_DIR]:
            os.makedirs(cache_dir, exist_ok=True)
        
        # Load caches and memory
        self._load_caches()
        self._load_memory()
        
        # Language mapping
        self.language_map = {
            'english': 'en', 'tiếng anh': 'en', 'anh': 'en',
            'chinese': 'zh', 'trung quốc': 'zh', 'trung': 'zh',
            'japanese': 'ja', 'nhật bản': 'ja', 'nhật': 'ja',
            'korean': 'ko', 'hàn quốc': 'ko', 'hàn': 'ko',
            'hangul': 'ko', 'tiếng hàn': 'ko',
            'french': 'fr', 'pháp': 'fr', 'german': 'de', 'đức': 'de',
            'spanish': 'es', 'tây ban nha': 'es',
            'vietnamese': 'vi', 'việt nam': 'vi', 'việt': 'vi'
        }
    
    def _load_caches(self):
        """Load tất cả cache files"""
        # File cache
        cache_file = os.path.join(CACHE_DIR, 'file_cache.json')
        if os.path.exists(cache_file):
            try:
                with open(cache_file, 'r', encoding='utf-8') as f:
                    self.file_cache = json.load(f)
            except: pass
        
        # Embedding cache
        embedding_file = os.path.join(EMBEDDING_CACHE_DIR, 'embedding_cache.pkl')
        metadata_file = os.path.join(EMBEDDING_CACHE_DIR, 'vector_metadata.json')
        if os.path.exists(embedding_file):
            try:
                with open(embedding_file, 'rb') as f:
                    self.embedding_cache = pickle.load(f)
            except: pass
        if os.path.exists(metadata_file):
            try:
                with open(metadata_file, 'r', encoding='utf-8') as f:
                    self.vector_metadata = json.load(f)
            except: pass
        
        # Translation cache
        translation_file = os.path.join(TRANSLATION_CACHE_DIR, 'translation_cache.json')
        if os.path.exists(translation_file):
            try:
                with open(translation_file, 'r', encoding='utf-8') as f:
                    self.translation_cache = json.load(f)
            except: pass
    
    def _load_memory(self):
        """Load memory and learning data"""
        memory_file = os.path.join(CACHE_DIR, 'memory.json')
        if os.path.exists(memory_file):
            try:
                with open(memory_file, 'r', encoding='utf-8') as f:
                    memory_data = json.load(f)
                    self.conversation_memory = memory_data.get('conversation_memory', [])
                    self.user_preferences.update(memory_data.get('user_preferences', {}))
                    self.learning_data.update(memory_data.get('learning_data', {}))
                logger.info("Memory data loaded")
            except: pass
    
    def _save_caches(self):
        """Lưu tất cả cache files"""
        try:
            # File cache
            with open(os.path.join(CACHE_DIR, 'file_cache.json'), 'w', encoding='utf-8') as f:
                json.dump(self.file_cache, f, ensure_ascii=False, indent=2)
            
            # Embedding cache
            with open(os.path.join(EMBEDDING_CACHE_DIR, 'embedding_cache.pkl'), 'wb') as f:
                pickle.dump(self.embedding_cache, f)
            with open(os.path.join(EMBEDDING_CACHE_DIR, 'vector_metadata.json'), 'w', encoding='utf-8') as f:
                json.dump(self.vector_metadata, f, ensure_ascii=False, indent=2)
            
            # Translation cache
            with open(os.path.join(TRANSLATION_CACHE_DIR, 'translation_cache.json'), 'w', encoding='utf-8') as f:
                json.dump(self.translation_cache, f, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.error(f"Error saving caches: {e}")
    
    def _save_memory(self):
        """Save memory and learning data"""
        try:
            memory_data = {
                'conversation_memory': self.conversation_memory[-50:],  # Keep last 50 conversations
                'user_preferences': self.user_preferences,
                'learning_data': self.learning_data
            }
            with open(os.path.join(CACHE_DIR, 'memory.json'), 'w', encoding='utf-8') as f:
                json.dump(memory_data, f, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.error(f"Error saving memory: {e}")
    
    def _get_hash(self, text: str) -> str:
        """Tạo hash cho text"""
        return hashlib.md5(text.encode('utf-8')).hexdigest()
    
    def learn_from_interaction(self, query: str, response: str, feedback: str = None):
        """Learn from user interactions"""
        interaction = {
            'timestamp': datetime.now().isoformat(),
            'query': query,
            'response_length': len(response),
            'query_type': self._classify_query_type(query),
            'feedback': feedback,
            'source': self.current_source
        }
        
        # Add to conversation memory
        self.conversation_memory.append(interaction)
        
        # Update user preferences
        self._update_user_preferences(query, response, feedback)
        
        # Track successful/failed queries
        if feedback == 'positive':
            self.learning_data['successful_queries'].append(query)
        elif feedback == 'negative':
            self.learning_data['failed_queries'].append(query)
        
        # Save memory
        self._save_memory()
    
    def _classify_query_type(self, query: str) -> str:
        """Classify query type for learning"""
        query_lower = query.lower()
        
        if any(word in query_lower for word in ['tóm tắt', 'summary', 'tổng hợp']):
            return 'summary'
        elif any(word in query_lower for word in ['so sánh', 'compare', 'khác nhau']):
            return 'comparison'
        elif any(word in query_lower for word in ['dịch', 'translate']):
            return 'translation'
        elif any(word in query_lower for word in ['tìm', 'find', 'search', 'thông tin']):
            return 'search'
        elif '?' in query:
            return 'question'
        else:
            return 'general'
    
    def _update_user_preferences(self, query: str, response: str, feedback: str):
        """Update user preferences based on interactions"""
        # Track preferred answer length
        response_length = len(response)
        if feedback == 'positive':
            if response_length < 200:
                self.user_preferences['preferred_answer_length'] = 'short'
            elif response_length > 800:
                self.user_preferences['preferred_answer_length'] = 'long'
            else:
                self.user_preferences['preferred_answer_length'] = 'medium'
        
        # Track frequently asked topics
        query_type = self._classify_query_type(query)
        if query_type in self.user_preferences['frequently_asked_topics']:
            self.user_preferences['frequently_asked_topics'][query_type] += 1
        else:
            self.user_preferences['frequently_asked_topics'][query_type] = 1
    
    def get_conversation_context(self, current_query: str) -> str:
        """Get context from previous conversations"""
        if not self.conversation_memory:
            return ""
        
        # Get last 3 conversations for context
        recent_conversations = self.conversation_memory[-3:]
        context_parts = []
        
        for conv in recent_conversations:
            context_parts.append(f"Previous Q: {conv['query'][:100]}...")
        
        return "\n".join(context_parts)
    
    def set_data_source(self, source: str):
        """Set current data source: 'drive' or 'upload'"""
        self.current_source = source
    
    def detect_language(self, text: str) -> str:
        """Detect ngôn ngữ sử dụng OpenAI GPT-4.1-mini"""
        try:
            response = self.client.chat.completions.create(
                model="gpt-4.1-mini",
                messages=[
                    {"role": "system", "content": "Detect the language of the given text and return only the language code (en/vi/zh/ja/ko/fr/de/es)."},
                    {"role": "user", "content": f"Text: {text[:100]}"}
                ],
                max_tokens=10,
                temperature=0
            )
            lang = response.choices[0].message.content.strip().lower()
            return lang if len(lang) == 2 else 'vi'
        except:
            return 'vi'
    
    def translate_text(self, text: str, target_language: str) -> str:
        """Dịch text với OpenAI GPT-4.1-mini và caching"""
        cache_key = f"{self._get_hash(text)}_{target_language}"
        
        if cache_key in self.translation_cache:
            return self.translation_cache[cache_key]
        
        try:
            lang_names = {
                'en': 'English', 'vi': 'Vietnamese', 'zh': 'Chinese', 
                'ja': 'Japanese', 'ko': 'Korean', 'fr': 'French', 
                'de': 'German', 'es': 'Spanish'
            }
            target_name = lang_names.get(target_language, target_language)
            
            response = self.client.chat.completions.create(
                model="gpt-4.1-mini",
                messages=[
                    {"role": "system", "content": f"Translate the following text to {target_name}. Keep the original formatting and structure."},
                    {"role": "user", "content": text}
                ],
                temperature=0.3
            )
            
            translated = response.choices[0].message.content.strip()
            
            self.translation_cache[cache_key] = translated
            self._save_caches()
            return translated
        except Exception as e:
            return f"❌ Lỗi dịch: {str(e)}"
    
    def _detect_translation_request(self, query: str) -> Optional[Dict[str, str]]:
        """Detect translation request"""
        query_lower = query.lower().strip()
        patterns = [
            r'dịch\s+sang\s+(.+?)(?:\s|$)',
            r'translate\s+to\s+(.+?)(?:\s|$)',
            r'dịch\s+(.+?)(?:\s|$)',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, query_lower)
            if match:
                target_lang_text = match.group(1).strip()
                for lang_name, lang_code in self.language_map.items():
                    if lang_name in target_lang_text:
                        return {
                            'action': 'translate',
                            'target_language': lang_code,
                            'target_language_name': lang_name.title()
                        }
                return {'action': 'translate', 'target_language': 'en', 'target_language_name': 'English'}
        return None
    
    def authenticate_google_drive(self) -> bool:
        """Xác thực Google Drive"""
        try:
            creds = None
            if os.path.exists(TOKEN_FILE):
                with open(TOKEN_FILE, 'rb') as token:
                    creds = pickle.load(token)
            
            if not creds or not creds.valid:
                if creds and creds.expired and creds.refresh_token:
                    creds.refresh(Request())
                else:
                    if not os.path.exists(CREDENTIALS_FILE):
                        st.error("❌ Không tìm thấy credentials.json")
                        return False
                    flow = InstalledAppFlow.from_client_secrets_file(CREDENTIALS_FILE, SCOPES)
                    creds = flow.run_local_server(port=0)
                
                with open(TOKEN_FILE, 'wb') as token:
                    pickle.dump(creds, token)
            
            self.drive_service = build('drive', 'v3', credentials=creds)
            self.drive_service.about().get(fields="user").execute()
            return True
        except Exception as e:
            st.error(f"❌ Lỗi xác thực: {e}")
            return False
    
    def get_files_from_drive(self, folder_id: Optional[str] = None, 
                           file_types: Optional[List[str]] = None, max_files: int = 100) -> List[Dict]:
        """Lấy files từ Google Drive"""
        if not self.drive_service:
            return []
        
        try:
            query_parts = ["trashed=false"]
            if folder_id:
                query_parts.append(f"parents in '{folder_id}'")
            
            if file_types:
                mime_map = {
                    'pdf': 'application/pdf',
                    'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    'txt': 'text/plain',
                    'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
                }
                mime_queries = [f"mimeType='{mime_map[ft]}'" for ft in file_types if ft in mime_map]
                if mime_queries:
                    query_parts.append(f"({' or '.join(mime_queries)})")
            
            query = " and ".join(query_parts)
            results = self.drive_service.files().list(
                q=query, pageSize=max_files,
                fields="files(id, name, mimeType, size, modifiedTime, webViewLink)"
            ).execute()
            
            files = results.get('files', [])
            filtered_files = []
            for file_info in files:
                size = int(file_info.get('size', 0))
                if size <= self.max_file_size_mb * 1024 * 1024:
                    filtered_files.append(file_info)
            
            return filtered_files[:max_files]
        except Exception as e:
            st.error(f"❌ Lỗi lấy files: {e}")
            return []
    
    def _extract_file_content(self, file_content: io.BytesIO, mime_type: str) -> str:
        """Extract content from different file types"""
        try:
            if mime_type == 'application/pdf':
                pdf_reader = PyPDF2.PdfReader(file_content)
                return "\n".join([page.extract_text() for page in pdf_reader.pages])
            elif 'wordprocessingml' in mime_type:
                doc = DocxDocument(file_content)
                return "\n".join([para.text for para in doc.paragraphs])
            elif mime_type == 'text/plain':
                return file_content.read().decode('utf-8', errors='ignore')
            elif 'spreadsheetml' in mime_type:
                df = pd.read_excel(file_content, sheet_name=None)
                return "\n".join([f"Sheet {name}:\n{data.to_string()}" for name, data in df.items()])
            return ""
        except Exception as e:
            logger.error(f"Content extraction error: {e}")
            return ""

    def download_file_content(self, file_id: str, mime_type: str, file_name: str) -> str:
        """Download và extract file content"""
        try:
            file_info = self.drive_service.files().get(fileId=file_id, fields="modifiedTime").execute()
            file_hash = self._get_hash(f"{file_id}_{file_info['modifiedTime']}")
            
            if file_hash in self.file_cache:
                return self.file_cache[file_hash]
            
            request = self.drive_service.files().get_media(fileId=file_id)
            file_content = io.BytesIO()
            
            import googleapiclient.http
            downloader = googleapiclient.http.MediaIoBaseDownload(file_content, request)
            done = False
            while not done:
                status, done = downloader.next_chunk()
            
            file_content.seek(0)
            content = self._extract_file_content(file_content, mime_type)
            
            if content:
                self.file_cache[file_hash] = content
                self._save_caches()
            
            return content
        except Exception as e:
            return f"Lỗi tải file: {str(e)}"
    
    def process_documents(self, files_info: List[Dict]) -> List[Document]:
        """Xử lý documents từ Google Drive"""
        documents = []
        progress_bar = st.progress(0)
        
        for i, file_info in enumerate(files_info):
            progress_bar.progress((i + 1) / len(files_info))
            content = self.download_file_content(file_info['id'], file_info['mimeType'], file_info['name'])
            
            if content and len(content.strip()) > 50:
                doc = Document(
                    page_content=content,
                    metadata={
                        'source': f"[DRIVE] {file_info['name']}",
                        'file_id': file_info['id'],
                        'file_type': 'google_drive',
                        'web_view_link': file_info.get('webViewLink', ''),
                        'content_length': len(content)
                    }
                )
                documents.append(doc)
        
        progress_bar.empty()
        return documents
    
    def process_uploaded_files(self, uploaded_files) -> List[Document]:
        """Process uploaded files"""
        documents = []
        progress_bar = st.progress(0)
        
        for i, uploaded_file in enumerate(uploaded_files):
            progress_bar.progress((i + 1) / len(uploaded_files))
            
            try:
                file_bytes = uploaded_file.getvalue()
                content = self._extract_file_content(io.BytesIO(file_bytes), uploaded_file.type)
                
                if content and len(content.strip()) > 50:
                    doc = Document(
                        page_content=content,
                        metadata={
                            'source': f"[UPLOAD] {uploaded_file.name}",
                            'file_type': 'uploaded',
                            'file_size': len(file_bytes),
                            'content_length': len(content)
                        }
                    )
                    documents.append(doc)
                
            except Exception as e:
                st.warning(f"⚠️ Could not process {uploaded_file.name}: {str(e)}")
        
        progress_bar.empty()
        return documents
    
    def create_drive_vector_store(self, documents: List[Document]):
        """Tạo vector store riêng cho Google Drive files"""
        if not documents:
            st.warning("⚠️ Không có Drive documents")
            return
        
        self.drive_vector_store = self._create_vector_store_internal(documents, "drive")
        st.success(f"✅ Drive vector store created with {len(documents)} documents")
    
    def create_upload_vector_store(self, documents: List[Document]):
        """Tạo vector store riêng cho uploaded files"""
        if not documents:
            st.warning("⚠️ Không có uploaded documents")
            return
        
        self.upload_vector_store = self._create_vector_store_internal(documents, "upload")
        st.success(f"✅ Upload vector store created with {len(documents)} documents")
    
    def _create_vector_store_internal(self, documents: List[Document], source_type: str):
        """Internal method to create vector store"""
        # Split documents
        texts = self.text_splitter.split_documents(documents)
        st.info(f"📝 Created {len(texts)} text chunks from {source_type} files")
        
        # Use cached embeddings
        cached_embeddings = []
        texts_to_embed = []
        
        for text_doc in texts:
            text_hash = self._get_hash(text_doc.page_content)
            if text_hash in self.embedding_cache:
                cached_embeddings.append((text_doc, self.embedding_cache[text_hash]))
            else:
                texts_to_embed.append((text_doc, text_hash))
        
        # Create new embeddings with BATCH PROCESSING
        if texts_to_embed:
            with st.spinner(f"🤖 Creating embeddings for {source_type} files..."):
                batch_size = 20  # Process 20 chunks at once
                progress_bar = st.progress(0)
                
                for i in range(0, len(texts_to_embed), batch_size):
                    batch = texts_to_embed[i:i + batch_size]
                    
                    # Prepare batch texts
                    batch_texts = [text_doc.page_content for text_doc, _ in batch]
                    
                    # Single API call for entire batch - Much faster with OpenAI!
                    batch_embeddings = self.embeddings.embed_documents(batch_texts)
                    
                    # Cache all batch results
                    for j, (text_doc, text_hash) in enumerate(batch):
                        self.embedding_cache[text_hash] = batch_embeddings[j]
                        cached_embeddings.append((text_doc, batch_embeddings[j]))
                    
                    # Update progress
                    progress = min((i + batch_size) / len(texts_to_embed), 1.0)
                    progress_bar.progress(progress)
                
                progress_bar.empty()
        
        # Create vector store
        if cached_embeddings:
            texts_list = [text_doc for text_doc, _ in cached_embeddings]
            try:
                embeddings_list = [emb for _, emb in cached_embeddings]
                embeddings_array = np.array(embeddings_list).astype('float32')
                
                import faiss
                from langchain_community.vectorstores.faiss import FAISS
                from langchain_community.docstore.in_memory import InMemoryDocstore
                
                index = faiss.IndexFlatL2(embeddings_array.shape[1])
                index.add(embeddings_array)
                
                docstore = InMemoryDocstore({str(i): doc for i, doc in enumerate(texts_list)})
                index_to_docstore_id = {i: str(i) for i in range(len(texts_list))}
                
                vector_store = FAISS(
                    embedding_function=self.embeddings.embed_query,
                    index=index,
                    docstore=docstore,
                    index_to_docstore_id=index_to_docstore_id
                )
                
                # Save cache
                self._save_caches()
                
                return vector_store
            except:
                return FAISS.from_documents(texts, self.embeddings)
        else:
            return FAISS.from_documents(texts, self.embeddings)
    
    def search_similar_documents(self, query: str, k: Optional[int] = None) -> List[Document]:
        """Tìm kiếm documents theo source được chọn"""
        if not k:
            k = self.max_search_results
        
        try:
            if self.current_source == "drive" and self.drive_vector_store:
                return self.drive_vector_store.similarity_search(query, k=k)
            elif self.current_source == "upload" and self.upload_vector_store:
                return self.upload_vector_store.similarity_search(query, k=k)
            else:
                return []
        except Exception as e:
            logger.error(f"Search error: {e}")
            return []
    
    def generate_answer(self, query: str, context_docs: List[Document], target_language: str = 'vi') -> Dict[str, Any]:
        """Tạo câu trả lời với OpenAI và memory context"""
        if not context_docs:
            return {
                'answer': "🤔 Không tìm thấy thông tin liên quan trong tài liệu.",
                'confidence': 0,
                'sources_used': 0
            }
        
        # Get conversation context for continuity
        conversation_context = self.get_conversation_context(query)
        
        # Create context from documents
        context_parts = []
        for i, doc in enumerate(context_docs):
            source = doc.metadata.get('source', f'Document {i+1}')
            content = doc.page_content[:500]
            context_parts.append(f"[Source {i+1}: {source}]\n{content}")
        
        context = "\n\n".join(context_parts)
        
        # Customize prompt based on user preferences
        answer_length_instruction = ""
        if self.user_preferences['preferred_answer_length'] == 'short':
            answer_length_instruction = "6. Trả lời ngắn gọn, tối đa 2-3 câu"
        elif self.user_preferences['preferred_answer_length'] == 'long':
            answer_length_instruction = "6. Trả lời chi tiết, đầy đủ thông tin"
        else:
            answer_length_instruction = "6. Trả lời vừa phải, cân bằng giữa ngắn gọn và đầy đủ"
        
        # Add conversation context if available
        context_instruction = ""
        if conversation_context:
            context_instruction = f"\n\nBỐI CẢNH CUỘC TRÒ CHUYỆN TRƯỚC:\n{conversation_context}"
        
        system_prompt = f"""
Bạn là AI assistant thông minh với khả năng ghi nhớ và học hỏi. Trả lời câu hỏi dựa trên tài liệu được cung cấp.

NGUYÊN TẮC:
1. Chỉ dùng thông tin từ tài liệu
2. Trả lời rõ ràng và chi tiết
3. Trích dẫn nguồn khi cần
4. Kết nối với bối cảnh cuộc trò chuyện trước nếu có
5. Học hỏi từ feedback để cải thiện
{answer_length_instruction}

NGUỒN DỮ LIỆU: {self.current_source.upper()}
{context_instruction}
"""
        
        user_prompt = f"""
TÀI LIỆU HIỆN TẠI:
{context}

CÂU HỎI: {query}
"""
        
        try:
            response = self.client.chat.completions.create(
                model="gpt-4.1-mini",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.7,
                max_tokens=1500
            )
            
            answer = response.choices[0].message.content
            
            # Learn from this interaction
            self.learn_from_interaction(query, answer)
            
            return {
                'answer': answer,
                'confidence': min(len(context_docs) * 0.2, 1.0),
                'sources_used': len(context_docs)
            }
        except Exception as e:
            return {
                'answer': f"❌ Lỗi: {str(e)}",
                'confidence': 0,
                'sources_used': 0
            }
    
    def chat(self, query: str) -> Dict[str, Any]:
        """Main chat function với memory và learning"""
        start_time = time.time()
        
        # Detect translation request
        translation_request = self._detect_translation_request(query)
        
        if translation_request:
            # Handle translation
            if hasattr(st.session_state, 'messages') and st.session_state.messages:
                last_message = None
                for msg in reversed(st.session_state.messages):
                    if msg["role"] == "assistant":
                        last_message = msg["content"]
                        break
                
                if last_message:
                    translated = self.translate_text(last_message, translation_request['target_language'])
                    return {
                        'answer': f"🌐 **{translation_request['target_language_name']}:**\n\n{translated}",
                        'sources': [],
                        'num_sources': 0,
                        'confidence': 0.9,
                        'processing_time': time.time() - start_time,
                        'is_translation': True
                    }
            
            return {
                'answer': "❌ Không có câu trả lời để dịch.",
                'sources': [],
                'num_sources': 0,
                'confidence': 0,
                'processing_time': time.time() - start_time
            }
        
        # Normal query processing
        query_language = self.detect_language(query)
        similar_docs = self.search_similar_documents(query)
        response_data = self.generate_answer(query, similar_docs, target_language=query_language)
        
        sources = []
        for doc in similar_docs:
            sources.append({
                'file_name': doc.metadata.get('source', 'Unknown'),
                'web_view_link': doc.metadata.get('web_view_link', ''),
                'content_preview': doc.page_content[:150] + "..." if len(doc.page_content) > 150 else doc.page_content
            })
        
        return {
            'answer': response_data['answer'],
            'sources': sources,
            'num_sources': len(sources),
            'confidence': response_data['confidence'],
            'processing_time': time.time() - start_time,
            'query_language': query_language,
            'is_translation': False
        }
    
    def get_cache_stats(self) -> Dict[str, Any]:
        """Cache statistics"""
        return {
            'file_cache_count': len(self.file_cache),
            'embedding_cache_count': len(self.embedding_cache),
            'translation_cache_count': len(self.translation_cache),
            'conversation_memory_count': len(self.conversation_memory),
            'vector_metadata': self.vector_metadata
        }
    
    def clear_all_caches(self):
        """Clear all caches"""
        try:
            self.file_cache.clear()
            self.embedding_cache.clear()
            self.translation_cache.clear()
            self.vector_metadata.clear()
            self.conversation_memory.clear()
            self.user_preferences = {
                'preferred_answer_length': 'medium',
                'preferred_language': 'vi',
                'frequently_asked_topics': {},
                'user_interaction_patterns': []
            }
            self.learning_data = {
                'successful_queries': [],
                'failed_queries': [],
                'user_feedback': {},
                'context_patterns': {}
            }
            
            import shutil
            for cache_dir in [CACHE_DIR, EMBEDDING_CACHE_DIR, TRANSLATION_CACHE_DIR, VECTOR_STORE_DIR]:
                if os.path.exists(cache_dir):
                    shutil.rmtree(cache_dir)
                    if cache_dir != VECTOR_STORE_DIR:
                        os.makedirs(cache_dir, exist_ok=True)
            return True
        except:
            return False

# Streamlit UI
def main():
    st.set_page_config(
        page_title="Enhanced RAG Chatbot with OpenAI",
        page_icon="🤖",
        layout="wide"
    )
    
    # Custom CSS
    st.markdown("""
    <style>
    .main-header {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 1rem;
        border-radius: 10px;
        text-align: center;
        margin-bottom: 2rem;
    }
    .source-card {
        background: white;
        padding: 1rem;
        border-left: 4px solid #667eea;
        margin: 0.5rem 0;
        border-radius: 5px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    .translation-indicator {
        background: #e8f5e8;
        border-left: 4px solid #4caf50;
        padding: 0.5rem;
        margin: 0.5rem 0;
        border-radius: 5px;
    }
    .memory-indicator {
        background: #f0f8ff;
        border-left: 4px solid #4169e1;
        padding: 0.5rem;
        margin: 0.5rem 0;
        border-radius: 5px;
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Header
    st.markdown("""
    <div class="main-header">
        <h1>🤖 Enhanced RAG Chatbot</h1>
        <p>AI Assistant with OpenAI GPT-4.1-mini & text-embedding-3-large</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Sidebar
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        # API Key status
        st.markdown("### 🔑 OpenAI API Status")
        if OPENAI_API_KEY:
            # Show partial key for security
            masked_key = OPENAI_API_KEY[:10] + "*" * 20 + OPENAI_API_KEY[-4:]
            st.success(f"✅ OpenAI API Key: {masked_key}")
        else:
            st.error("❌ Please set OPENAI_API_KEY in .env file")
            st.code("OPENAI_API_KEY=sk-proj-your-actual-key", language="bash")
        
        # Data source selection with memory
        st.markdown("### 📊 Data Sources")
        data_source = st.radio(
            "Choose data source:",
            options=["Google Drive", "Upload Files"],
            index=0,
            help="Select where to get your documents from"
        )
        
        # Show current source status
        if 'chatbot' in st.session_state and st.session_state.chatbot:
            current_source = st.session_state.chatbot.current_source
            if current_source == "drive":
                st.info("🗄️ Currently using: Google Drive files")
            elif current_source == "upload":
                st.info("📤 Currently using: Uploaded files")
        
        # Memory status
        if 'chatbot' in st.session_state and st.session_state.chatbot:
            memory_count = len(st.session_state.chatbot.conversation_memory)
            if memory_count > 0:
                st.success(f"🧠 Memory: {memory_count} conversations learned")
            else:
                st.info("🧠 Memory: Fresh start")
        
        # Google Drive settings
        if data_source == "Google Drive":
            with st.expander("🗄️ Google Drive Settings", expanded=True):
                folder_id = st.text_input(
                    "📁 Folder ID (optional)",
                    value=os.getenv('GOOGLE_DRIVE_FOLDER_ID', ''),
                    help="Leave empty to read all files"
                )
        else:
            folder_id = None
        
        # File upload settings
        if data_source == "Upload Files":
            with st.expander("📤 Upload Files", expanded=True):
                # Add clear button at top
                col1, col2 = st.columns([3, 1])
                with col2:
                    if st.button("🗑️ Clear", help="Clear all uploaded files"):
                        if 'file_uploader_key' not in st.session_state:
                            st.session_state.file_uploader_key = 0
                        st.session_state.file_uploader_key += 1
                        st.rerun()
                
                # File uploader with dynamic key
                uploader_key = f"file_uploader_{st.session_state.get('file_uploader_key', 0)}"
                
                uploaded_files = st.file_uploader(
                    "Choose files",
                    accept_multiple_files=True,
                    type=['pdf', 'docx', 'txt', 'xlsx'],
                    help="Upload PDF, DOCX, TXT, or XLSX files",
                    key=uploader_key
                )
                
                if uploaded_files:
                    st.success(f"📁 {len(uploaded_files)} file(s) uploaded")
                    
                    for i, file in enumerate(uploaded_files, 1):
                        file_size = len(file.getvalue()) / 1024 / 1024
                        st.write(f"{i}. 📄 **{file.name}** ({file_size:.1f}MB)")
                    
                    total_size = sum(len(file.getvalue()) for file in uploaded_files) / 1024 / 1024
                    st.caption(f"Total size: {total_size:.1f}MB")
                else:
                    st.info("No files uploaded yet")
        else:
            uploaded_files = None
        
        with st.expander("🔧 Advanced Settings"):
            max_files = st.slider("Max Files", 10, 200, 50)
            chunk_size = st.slider("Chunk Size", 500, 2000, 1000)
            max_search_results = st.slider("Search Results", 3, 10, 5)
        
        with st.expander("🌐 Translation Help"):
            st.markdown("""
            **Usage:**
            - "dịch sang tiếng Anh"
            - "translate to English"
            - "dịch sang Chinese"
            - "dịch sang tiếng Hàn"   
            """)
        
        file_types = st.multiselect(
            "📄 File Types",
            ['pdf', 'docx', 'txt', 'xlsx'],
            default=['pdf', 'docx', 'txt']
        )
        
        # Cache stats
        if 'chatbot' in st.session_state and st.session_state.get('documents_loaded'):
            st.markdown("### 📊 Cache Stats")
            chatbot = st.session_state.chatbot
            stats = chatbot.get_cache_stats()
            
            st.metric("💾 File Cache", stats['file_cache_count'])
            st.metric("🤖 Embeddings", stats['embedding_cache_count'])
            st.metric("🌐 Translations", stats['translation_cache_count'])
            st.metric("🧠 Conversations", stats['conversation_memory_count'])
            
            if st.button("🗑️ Clear All Cache"):
                if chatbot.clear_all_caches():
                    st.success("All cache cleared!")
                    st.rerun()
        
        if st.button("🔄 Reset App"):
            for key in list(st.session_state.keys()):
                del st.session_state[key]
            st.rerun()
    
    # Initialize session state
    if 'chatbot' not in st.session_state:
        st.session_state.chatbot = None
    if 'messages' not in st.session_state:
        st.session_state.messages = []
    if 'documents_loaded' not in st.session_state:
        st.session_state.documents_loaded = False
    if 'file_uploader_key' not in st.session_state:
        st.session_state.file_uploader_key = 0
    
    # Main content
    col1, col2 = st.columns([2, 1])
    
    with col1:
        if st.button("🚀 Connect & Load Documents", type="primary"):
            if not OPENAI_API_KEY:
                st.error("❌ Please set OPENAI_API_KEY in .env file")
            else:
                with st.spinner("🔄 Initializing..."):
                    try:
                        # Initialize chatbot
                        chatbot = GoogleDriveRAGChatbot()
                        chatbot.chunk_size = chunk_size
                        chatbot.max_search_results = max_search_results
                        
                        # Set data source
                        if data_source == "Google Drive":
                            chatbot.set_data_source("drive")
                        elif data_source == "Upload Files":
                            chatbot.set_data_source("upload")
                        
                        drive_docs = []
                        upload_docs = []
                        
                        # Process Google Drive files
                        if data_source == "Google Drive":
                            if chatbot.authenticate_google_drive():
                                st.success("✅ Google Drive connected!")
                                
                                files = chatbot.get_files_from_drive(
                                    folder_id,
                                    file_types,
                                    max_files
                                )
                                
                                if files:
                                    st.info(f"📁 Found {len(files)} Drive files")
                                    drive_docs = chatbot.process_documents(files)
                                    if drive_docs:
                                        chatbot.create_drive_vector_store(drive_docs)
                                else:
                                    st.warning("⚠️ No Drive files found")
                            else:
                                st.error("❌ Could not connect to Google Drive")
                                return
                        
                        # Process uploaded files
                        elif data_source == "Upload Files":
                            if uploaded_files:
                                st.info(f"📤 Processing {len(uploaded_files)} uploaded files")
                                upload_docs = chatbot.process_uploaded_files(uploaded_files)
                                if upload_docs:
                                    chatbot.create_upload_vector_store(upload_docs)
                            else:
                                st.error("❌ No files uploaded")
                                return
                        
                        # Check if we have any documents processed
                        total_docs = len(drive_docs) + len(upload_docs)
                        if total_docs > 0:
                            st.session_state.chatbot = chatbot
                            st.session_state.documents_loaded = True
                            
                            st.session_state.file_stats = {
                                'total_files': total_docs,
                                'drive_files': len(drive_docs),
                                'upload_files': len(upload_docs),
                                'total_documents': total_docs,
                                'total_chunks': sum(len(chatbot.text_splitter.split_text(doc.page_content)) for doc in drive_docs + upload_docs),
                                'data_source': data_source,
                                'memory_conversations': len(chatbot.conversation_memory)
                            }
                            
                            st.success("🎉 System ready with OpenAI & memory!")
                            
                            # Show stats
                            stats = st.session_state.file_stats
                            if data_source == "Google Drive":
                                col_stat1, col_stat2, col_stat3 = st.columns(3)
                                with col_stat1:
                                    st.metric("📁 Drive Files", stats['drive_files'])
                                with col_stat2:
                                    st.metric("📝 Chunks", stats['total_chunks'])
                                with col_stat3:
                                    st.metric("🧠 Memory", stats['memory_conversations'])
                            else:  # Upload Files
                                col_stat1, col_stat2, col_stat3 = st.columns(3)
                                with col_stat1:
                                    st.metric("📤 Upload Files", stats['upload_files'])
                                with col_stat2:
                                    st.metric("📝 Chunks", stats['total_chunks'])
                                with col_stat3:
                                    st.metric("🧠 Memory", stats['memory_conversations'])
                        else:
                            st.error("❌ No documents found to process")
                            
                    except Exception as e:
                        st.error(f"❌ Error: {str(e)}")
    
    with col2:
        st.markdown("### 🎯 System Status")
        
        if st.session_state.documents_loaded:
            st.success("🟢 Ready")
            if 'file_stats' in st.session_state:
                stats = st.session_state.file_stats
                source_emoji = "🗄️" if stats.get('data_source') == "Google Drive" else "📤"
                st.markdown(f"""
                **📊 Statistics:**
                - {source_emoji} Source: {stats.get('data_source', 'Unknown')}
                - Files: {stats.get('total_files', 0)}
                - Chunks: {stats.get('total_chunks', 0)}
                - 🧠 Memory: {stats.get('memory_conversations', 0)} conversations
                
                **🤖 AI Models:**
                - Chat: GPT-4.1-mini
                - Embeddings: text-embedding-3-large
                """)
                
                # Show learning insights
                if 'chatbot' in st.session_state and st.session_state.chatbot.user_preferences['frequently_asked_topics']:
                    topics = st.session_state.chatbot.user_preferences['frequently_asked_topics']
                    most_common = max(topics, key=topics.get)
                    st.caption(f"💡 Most asked: {most_common} ({topics[most_common]}x)")
        else:
            st.warning("🟡 Not connected")
            st.info("👆 Choose data source and click 'Connect & Load Documents'")
        
        with st.expander("💡 Usage Tips"):
            st.markdown("""
            **System Status:**
            - API: ✅ OpenAI GPT-4.1-mini
            - Embeddings: 🎯 text-embedding-3-large
            - Memory: 🧠 Learning enabled
            - Sources: 📊 Separated by type
            
            **Effective Questions:**
            - Ask specific questions about document content
            - Request summaries or analysis
            - Search for detailed information
            
            **Translation:**
            - "dịch sang tiếng Anh"
            - "translate to Chinese"
            - "dịch sang Japanese"
            
            **Memory Features:**
            - Remembers your conversation history
            - Learns your preferred answer style
            - Adapts to your question patterns
            """)
    
    # Chat interface - same as before but with OpenAI indicators
    if st.session_state.documents_loaded and st.session_state.chatbot:
        st.markdown("---")
        st.header("💬 Chat with OpenAI Assistant")
        
        # Display chat history
        for i, message in enumerate(st.session_state.messages):
            with st.chat_message(message["role"]):
                st.write(message["content"])
                
                # Show metadata for assistant messages
                if message["role"] == "assistant" and "metadata" in message:
                    metadata = message["metadata"]
                    
                    # Translation indicator
                    if metadata.get('is_translation'):
                        st.markdown("""
                        <div class="translation-indicator">
                            🌐 This is an automatic translation using OpenAI GPT-4.1-mini
                        </div>
                        """, unsafe_allow_html=True)
                    
                    # Memory context indicator
                    if i > 0:  # Not the first message
                        st.markdown("""
                        <div class="memory-indicator">
                            🧠 Using conversation context and learned preferences (GPT-4.1-mini powered)
                        </div>
                        """, unsafe_allow_html=True)
                    
                    # Show metrics
                    col_meta1, col_meta2, col_meta3, col_meta4 = st.columns(4)
                    with col_meta1:
                        confidence = metadata.get('confidence', 0)
                        st.caption(f"🎯 Confidence: {confidence:.1%}")
                    with col_meta2:
                        processing_time = metadata.get('processing_time', 0)
                        st.caption(f"⏱️ Time: {processing_time:.1f}s")
                    with col_meta3:
                        num_sources = metadata.get('num_sources', 0)
                        st.caption(f"📚 Sources: {num_sources}")
                    with col_meta4:
                        source_icon = "🗄️" if st.session_state.chatbot.current_source == "drive" else "📤"
                        st.caption(f"{source_icon} OpenAI")
                    
                    # Show sources
                    if metadata.get('sources') and not metadata.get('is_translation'):
                        with st.expander(f"📚 Sources ({len(metadata['sources'])} documents)", expanded=False):
                            for j, source in enumerate(metadata['sources'], 1):
                                source_type = "📤 Uploaded" if '[UPLOAD]' in source['file_name'] else "🗄️ Google Drive"
                                st.markdown(f"""
                                <div class="source-card">
                                    <strong>{j}. {source['file_name']}</strong>
                                    {f'<br><a href="{source["web_view_link"]}" target="_blank">🔗 View file</a>' if source.get('web_view_link') else ''}
                                    <br><em>{source['content_preview']}</em>
                                    <br><small>Source: {source_type} | Powered by OpenAI</small>
                                </div>
                                """, unsafe_allow_html=True)
        
        # Input area
        st.markdown("### 💭 Ask a Question")
        
        # Quick suggestions
        if not st.session_state.messages:
            st.markdown("**💡 Quick suggestions:**")
            suggestion_cols = st.columns(4)
            
            suggestions = [
                "Summarize main content",
                "Key points",
                "Find information about...",
                "dịch sang English"
            ]
            
            for i, suggestion in enumerate(suggestions):
                with suggestion_cols[i]:
                    if st.button(suggestion, key=f"suggestion_{i}"):
                        st.session_state.auto_query = suggestion
        
        # Chat input
        query = st.chat_input(
            "Type your question or 'dịch sang [language]'... (Powered by GPT-4.1-mini)",
            key="chat_input"
        )
        
        # Handle auto query from suggestions
        if 'auto_query' in st.session_state:
            query = st.session_state.auto_query
            del st.session_state.auto_query
        
        if query:
            # Validate query
            if len(query.strip()) < 2:
                st.warning("⚠️ Question too short. Please provide more details.")
                st.stop()
            
            # Display user message
            with st.chat_message("user"):
                st.write(query)
            
            # Add to history
            st.session_state.messages.append({"role": "user", "content": query})
            
            # Generate response
            with st.chat_message("assistant"):
                with st.spinner("🤔 GPT-4.1-mini is thinking..."):
                    try:
                        response = st.session_state.chatbot.chat(query)
                        
                        # Display response
                        st.write(response['answer'])
                        
                        # Translation indicator
                        if response.get('is_translation'):
                            st.markdown("""
                            <div class="translation-indicator">
                                🌐 This is an automatic translation using OpenAI GPT-4.1-mini
                            </div>
                            """, unsafe_allow_html=True)
                        
                        # Memory context indicator
                        if len(st.session_state.messages) > 1:  # Not the first conversation
                            st.markdown("""
                            <div class="memory-indicator">
                                🧠 Using conversation context and learned preferences (GPT-4.1-mini powered)
                            </div>
                            """, unsafe_allow_html=True)
                        
                        # Display metadata
                        col_meta1, col_meta2, col_meta3, col_meta4 = st.columns(4)
                        with col_meta1:
                            st.caption(f"🎯 Confidence: {response['confidence']:.1%}")
                        with col_meta2:
                            st.caption(f"⏱️ Time: {response['processing_time']:.1f}s")
                        with col_meta3:
                            st.caption(f"📚 Sources: {response['num_sources']}")
                        with col_meta4:
                            source_icon = "🗄️" if st.session_state.chatbot.current_source == "drive" else "📤"
                            st.caption(f"{source_icon} GPT-4.1")
                        
                        # Display sources
                        if response['sources'] and not response.get('is_translation'):
                            with st.expander(f"📚 Sources ({response['num_sources']} documents)", expanded=False):
                                for j, source in enumerate(response['sources'], 1):
                                    source_type = "📤 Uploaded" if '[UPLOAD]' in source['file_name'] else "🗄️ Google Drive"
                                    st.markdown(f"""
                                    <div class="source-card">
                                        <strong>{j}. {source['file_name']}</strong>
                                        {f'<br><a href="{source["web_view_link"]}" target="_blank">🔗 View file</a>' if source.get('web_view_link') else ''}
                                        <br><em>{source['content_preview']}</em>
                                        <br><small>Source: {source_type} | Powered by GPT-4.1-mini</small>
                                    </div>
                                    """, unsafe_allow_html=True)
                        
                        # Learning feedback section
                        st.markdown("---")
                        st.markdown("**🧠 Help GPT-4.1-mini learn:**")
                        feedback_col1, feedback_col2, feedback_col3 = st.columns([1, 1, 3])
                        
                        with feedback_col1:
                            if st.button("👍 Good", key=f"thumbs_up_{len(st.session_state.messages)}"):
                                st.session_state.chatbot.learn_from_interaction(query, response['answer'], 'positive')
                                st.success("Thanks! GPT-4.1-mini will remember this works well.")
                        
                        with feedback_col2:
                            if st.button("👎 Poor", key=f"thumbs_down_{len(st.session_state.messages)}"):
                                st.session_state.chatbot.learn_from_interaction(query, response['answer'], 'negative')
                                st.info("Got it! GPT-4.1-mini will try to improve.")
                        
                        with feedback_col3:
                            st.caption("Your feedback helps GPT-4.1-mini learn your preferences")
                        
                    except Exception as e:
                        st.error(f"❌ Error processing question: {str(e)}")
                        response = {
                            'answer': "Sorry, an error occurred while processing your question.",
                            'sources': [],
                            'num_sources': 0,
                            'confidence': 0,
                            'processing_time': 0,
                            'is_translation': False
                        }
            
            # Add to history with metadata
            st.session_state.messages.append({
                "role": "assistant",
                "content": response['answer'],
                "metadata": {
                    'sources': response['sources'],
                    'num_sources': response['num_sources'],
                    'confidence': response['confidence'],
                    'processing_time': response['processing_time'],
                    'is_translation': response.get('is_translation', False)
                }
            })
    
    else:
        # Setup instructions when not connected
        st.markdown("---")
        st.info("🚀 **Connect to your data source to get started with GPT-4.1-mini!**")
        
        # Setup checklist
        st.markdown("### ✅ Setup Checklist:")
        
        checklist_items = [
            ("🔑 OpenAI API Key configured", bool(FIXED_OPENAI_API_KEY and FIXED_OPENAI_API_KEY != "your-openai-api-key-here")),
            ("📄 File types selected", bool(file_types)),
        ]
        
        # Add specific checks based on data source
        if data_source == "Google Drive":
            checklist_items.append(("📁 credentials.json file", os.path.exists(CREDENTIALS_FILE)))
        if data_source == "Upload Files":
            has_uploads = uploaded_files is not None and len(uploaded_files) > 0
            checklist_items.append(("📤 Files uploaded", has_uploads))
        
        for item, status in checklist_items:
            if status:
                st.success(f"✅ {item}")
            else:
                st.error(f"❌ {item}")
        
        # Feature preview
        if all(status for _, status in checklist_items):
            st.success("🎉 You're ready! Click 'Connect & Load Documents' above.")
            
            # Show enhanced features
            st.markdown("### 🆕 Enhanced Features with OpenAI:")
            col_feat1, col_feat2 = st.columns(2)
            
            with col_feat1:
                st.markdown("""
                **🤖 OpenAI GPT-4.1-mini:**
                - Latest generation model
                - Enhanced reasoning capabilities
                - Better context understanding
                
                **🎯 text-embedding-3-large:**
                - High-quality embeddings
                - Better semantic search
                - Improved relevance
                """)
            
            with col_feat2:
                st.markdown("""
                **🧠 Memory & Learning:**
                - Remembers conversation history
                - Learns your preferences
                - Adapts answer style
                
                **🌐 Translation Support:**
                - Multiple language support
                - Auto language detection
                - Translation caching
                
                **💾 Smart Caching:**
                - Intelligent embedding storage
                - Separate vector stores by source
                - Faster loading times
                """)
        else:
            st.warning("⚠️ Please complete the checklist before proceeding.")

if __name__ == "__main__":
    main()
