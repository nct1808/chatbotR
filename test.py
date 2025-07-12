import os
import json
import pickle
import logging
from typing import List, Dict, Any, Optional
from datetime import datetime
import streamlit as st
from dotenv import load_dotenv
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
import google.generativeai as genai
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import FAISS
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain.docstore.document import Document
import io
from docx import Document as DocxDocument
import PyPDF2
import pandas as pd
import hashlib
import time

# Load environment variables
load_dotenv()

# Cấu hình logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Constants
SCOPES = ['https://www.googleapis.com/auth/drive.readonly']
CREDENTIALS_FILE = 'credentials.json'
TOKEN_FILE = 'token.pickle'
VECTOR_STORE_DIR = 'vector_store'
CACHE_DIR = 'cache'

class GoogleDriveRAGChatbot:
    def __init__(self, gemini_api_key: Optional[str] = None):
        """
        Khởi tạo chatbot RAG với cấu hình nâng cao
        """
        # Lấy API key từ env hoặc parameter
        self.gemini_api_key = gemini_api_key or os.getenv('GEMINI_API_KEY')
        
        if not self.gemini_api_key:
            raise ValueError("Cần có Gemini API Key")
        
        # Cấu hình Gemini
        genai.configure(api_key=self.gemini_api_key)
        self.model = genai.GenerativeModel('gemini-1.5-flash-8b')
        
        # Khởi tạo embeddings
        self.embeddings = GoogleGenerativeAIEmbeddings(
            model="models/embedding-001",
            google_api_key=self.gemini_api_key
        )
        
        # Cấu hình từ environment
        self.chunk_size = int(os.getenv('CHUNK_SIZE', 1000))
        self.chunk_overlap = int(os.getenv('CHUNK_OVERLAP', 200))
        self.max_search_results = int(os.getenv('MAX_SEARCH_RESULTS', 5))
        self.max_file_size_mb = int(os.getenv('MAX_FILE_SIZE_MB', 50))
        
        # Text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
        )
        
        # Vector store và cache
        self.vector_store = None
        self.drive_service = None
        self.file_cache = {}
        
        # Tạo thư mục cache nếu chưa có
        os.makedirs(CACHE_DIR, exist_ok=True)
        
        # Load cache nếu có
        self._load_file_cache()
    
    def _load_file_cache(self):
        """Load cache của files đã xử lý"""
        cache_file = os.path.join(CACHE_DIR, 'file_cache.json')
        if os.path.exists(cache_file):
            try:
                with open(cache_file, 'r', encoding='utf-8') as f:
                    self.file_cache = json.load(f)
                logger.info(f"Loaded cache for {len(self.file_cache)} files")
            except Exception as e:
                logger.error(f"Error loading cache: {e}")
                self.file_cache = {}
    
    def _save_file_cache(self):
        """Lưu cache của files"""
        cache_file = os.path.join(CACHE_DIR, 'file_cache.json')
        try:
            with open(cache_file, 'w', encoding='utf-8') as f:
                json.dump(self.file_cache, f, ensure_ascii=False, indent=2)
            logger.info("File cache saved")
        except Exception as e:
            logger.error(f"Error saving cache: {e}")
    
    def _get_file_hash(self, file_id: str, modified_time: str) -> str:
        """Tạo hash để check file đã thay đổi chưa"""
        return hashlib.md5(f"{file_id}_{modified_time}".encode()).hexdigest()
    
    def authenticate_google_drive(self) -> bool:
        """Xác thực Google Drive API với error handling tốt hơn"""
        try:
            creds = None
            
            # Kiểm tra token đã lưu
            if os.path.exists(TOKEN_FILE):
                with open(TOKEN_FILE, 'rb') as token:
                    creds = pickle.load(token)
            
            # Refresh token nếu cần
            if not creds or not creds.valid:
                if creds and creds.expired and creds.refresh_token:
                    creds.refresh(Request())
                else:
                    if not os.path.exists(CREDENTIALS_FILE):
                        st.error("❌ Không tìm thấy file credentials.json")
                        st.info("📋 Hướng dẫn tạo credentials.json:")
                        st.markdown("""
                        1. Truy cập [Google Cloud Console](https://console.cloud.google.com/)
                        2. Tạo project và enable Google Drive API
                        3. Tạo OAuth 2.0 credentials (Desktop application)
                        4. Tải file JSON và đổi tên thành credentials.json
                        """)
                        return False
                    
                    flow = InstalledAppFlow.from_client_secrets_file(
                        CREDENTIALS_FILE, SCOPES)
                    creds = flow.run_local_server(port=0)
                
                # Lưu credentials
                with open(TOKEN_FILE, 'wb') as token:
                    pickle.dump(creds, token)
            
            # Tạo service
            self.drive_service = build('drive', 'v3', credentials=creds)
            
            # Test connection
            self.drive_service.about().get(fields="user").execute()
            logger.info("Google Drive authentication successful")
            return True
            
        except Exception as error:
            logger.error(f"Authentication error: {error}")
            st.error(f"❌ Lỗi xác thực: {error}")
            return False
    
    def get_files_from_drive(self, 
                           folder_id: Optional[str] = None, 
                           file_types: Optional[List[str]] = None,
                           max_files: int = 100) -> List[Dict]:
        """
        Lấy danh sách files từ Google Drive với filtering nâng cao
        """
        if not self.drive_service:
            return []
        
        try:
            # Xây dựng query
            query_parts = ["trashed=false"]
            
            if folder_id:
                query_parts.append(f"parents in '{folder_id}'")
            
            # Filter theo file types
            if file_types:
                mime_queries = []
                mime_map = {
                    'pdf': 'application/pdf',
                    'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    'txt': 'text/plain',
                    'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
                }
                
                for file_type in file_types:
                    if file_type in mime_map:
                        mime_queries.append(f"mimeType='{mime_map[file_type]}'")
                
                if mime_queries:
                    query_parts.append(f"({' or '.join(mime_queries)})")
            
            # Filter theo kích thước (Google Drive API không hỗ trợ trực tiếp)
            query = " and ".join(query_parts)
            
            # Lấy files với pagination
            all_files = []
            page_token = None
            
            while len(all_files) < max_files:
                results = self.drive_service.files().list(
                    q=query,
                    pageSize=min(100, max_files - len(all_files)),
                    pageToken=page_token,
                    fields="nextPageToken, files(id, name, mimeType, size, modifiedTime, webViewLink)"
                ).execute()
                
                files = results.get('files', [])
                
                # Filter theo kích thước
                filtered_files = []
                for file_info in files:
                    file_size = int(file_info.get('size', 0))
                    if file_size <= self.max_file_size_mb * 1024 * 1024:
                        filtered_files.append(file_info)
                    else:
                        logger.warning(f"File {file_info['name']} quá lớn ({file_size/1024/1024:.1f}MB)")
                
                all_files.extend(filtered_files)
                
                page_token = results.get('nextPageToken')
                if not page_token:
                    break
            
            logger.info(f"Found {len(all_files)} eligible files")
            return all_files[:max_files]
            
        except HttpError as error:
            logger.error(f"Error getting files: {error}")
            st.error(f"❌ Lỗi khi lấy danh sách files: {error}")
            return []
    
    def download_file_content(self, file_id: str, mime_type: str, file_name: str) -> str:
        """
        Tải nội dung file với caching và error handling
        """
        try:
            # Kiểm tra cache trước
            file_info = self.drive_service.files().get(
                fileId=file_id, 
                fields="modifiedTime"
            ).execute()
            
            file_hash = self._get_file_hash(file_id, file_info['modifiedTime'])
            
            # Kiểm tra cache
            if file_hash in self.file_cache:
                logger.info(f"Using cached content for {file_name}")
                return self.file_cache[file_hash]
            
            # Tải file mới
            logger.info(f"Downloading {file_name}")
            request = self.drive_service.files().get_media(fileId=file_id)
            file_content = io.BytesIO()
            
            import googleapiclient.http
            downloader = googleapiclient.http.MediaIoBaseDownload(file_content, request)
            done = False
            while done is False:
                status, done = downloader.next_chunk()
                if status:
                    progress = int(status.progress() * 100)
                    if progress % 20 == 0:  # Log mỗi 20%
                        logger.info(f"Download progress: {progress}%")
            
            file_content.seek(0)
            
            # Xử lý theo loại file
            content = ""
            if mime_type == 'application/pdf':
                content = self._extract_pdf_text(file_content)
            elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                content = self._extract_docx_text(file_content)
            elif mime_type == 'text/plain':
                content = file_content.read().decode('utf-8', errors='ignore')
            elif mime_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
                content = self._extract_xlsx_text(file_content)
            else:
                logger.warning(f"Unsupported file type: {mime_type}")
                return "Loại file không được hỗ trợ"
            
            # Lưu vào cache
            if content:
                self.file_cache[file_hash] = content
                self._save_file_cache()
            
            return content
                
        except Exception as error:
            logger.error(f"Error downloading file {file_name}: {error}")
            return f"Lỗi khi tải file: {str(error)}"
    
    def _extract_pdf_text(self, file_content: io.BytesIO) -> str:
        """Trích xuất text từ PDF với error handling"""
        try:
            pdf_reader = PyPDF2.PdfReader(file_content)
            text = ""
            for i, page in enumerate(pdf_reader.pages):
                try:
                    text += page.extract_text() + "\n"
                except Exception as e:
                    logger.warning(f"Error extracting page {i}: {e}")
                    continue
            return text.strip()
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            return f"Lỗi khi đọc PDF: {str(e)}"
    
    def _extract_docx_text(self, file_content: io.BytesIO) -> str:
        """Trích xuất text từ DOCX"""
        try:
            doc = DocxDocument(file_content)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Xử lý tables nếu có
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            return f"Lỗi khi đọc DOCX: {str(e)}"
    
    def _extract_xlsx_text(self, file_content: io.BytesIO) -> str:
        """Trích xuất text từ XLSX"""
        try:
            df = pd.read_excel(file_content, sheet_name=None)  # Đọc tất cả sheets
            text = ""
            
            for sheet_name, sheet_data in df.items():
                text += f"\n=== Sheet: {sheet_name} ===\n"
                text += sheet_data.to_string(index=False) + "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"XLSX extraction error: {e}")
            return f"Lỗi khi đọc XLSX: {str(e)}"
    
    def process_documents(self, files_info: List[Dict]) -> List[Document]:
        """
        Xử lý documents với progress tracking
        """
        documents = []
        
        # Progress bar
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        for i, file_info in enumerate(files_info):
            # Update progress
            progress = (i + 1) / len(files_info)
            progress_bar.progress(progress)
            status_text.text(f"Đang xử lý: {file_info['name']} ({i+1}/{len(files_info)})")
            
            # Tải nội dung file
            content = self.download_file_content(
                file_info['id'], 
                file_info['mimeType'], 
                file_info['name']
            )
            
            if content and len(content.strip()) > 50:  # Chỉ xử lý file có nội dung
                # Tạo Document object
                doc = Document(
                    page_content=content,
                    metadata={
                        'source': file_info['name'],
                        'file_id': file_info['id'],
                        'mime_type': file_info['mimeType'],
                        'size': file_info.get('size', 0),
                        'modified_time': file_info.get('modifiedTime', ''),
                        'web_view_link': file_info.get('webViewLink', ''),
                        'content_length': len(content)
                    }
                )
                documents.append(doc)
            else:
                logger.warning(f"File {file_info['name']} has no extractable content")
        
        # Clear progress
        progress_bar.empty()
        status_text.empty()
        
        logger.info(f"Processed {len(documents)} documents successfully")
        return documents
    
    def create_vector_store(self, documents: List[Document]):
        """
        Tạo vector store với caching
        """
        if not documents:
            st.warning("⚠️ Không có documents để xử lý")
            return
        
        # Kiểm tra xem đã có vector store cache chưa
        if os.path.exists(VECTOR_STORE_DIR):
            try:
                self.vector_store = FAISS.load_local(VECTOR_STORE_DIR, self.embeddings)
                st.info("📚 Đã tải vector store từ cache")
                return
            except Exception as e:
                logger.warning(f"Could not load cached vector store: {e}")
        
        # Chia nhỏ documents
        with st.spinner("🔧 Đang chia nhỏ documents..."):
            texts = self.text_splitter.split_documents(documents)
        
        st.info(f"📝 Đã tạo {len(texts)} đoạn text từ {len(documents)} documents")
        
        # Tạo embeddings với progress
        with st.spinner("🤖 Đang tạo embeddings..."):
            # Batch processing để tránh rate limit
            batch_size = 10
            all_texts = []
            
            for i in range(0, len(texts), batch_size):
                batch = texts[i:i+batch_size]
                all_texts.extend(batch)
                
                if i > 0:  # Delay để tránh rate limit
                    time.sleep(1)
        
        # Tạo vector store
        self.vector_store = FAISS.from_documents(all_texts, self.embeddings)
        
        # Lưu cache
        try:
            self.vector_store.save_local(VECTOR_STORE_DIR)
            logger.info("Vector store saved to cache")
        except Exception as e:
            logger.warning(f"Could not save vector store cache: {e}")
        
        st.success(f"✅ Đã tạo vector store với {len(all_texts)} đoạn text")
    
    def search_similar_documents(self, query: str, k: Optional[int] = None) -> List[Document]:
        """
        Tìm kiếm documents tương tự với scoring
        """
        if not self.vector_store:
            return []
        
        k = k or self.max_search_results
        
        try:
            # Tìm kiếm với score
            similar_docs_with_scores = self.vector_store.similarity_search_with_score(query, k=k)
            
            # Filter theo threshold score (tùy chọn)
            score_threshold = 0.8
            filtered_docs = [
                doc for doc, score in similar_docs_with_scores 
                if score < score_threshold  # FAISS trả về distance (càng nhỏ càng tốt)
            ]
            
            if not filtered_docs:
                # Nếu không có doc nào đạt threshold, lấy top results
                filtered_docs = [doc for doc, _ in similar_docs_with_scores[:k//2]]
            
            logger.info(f"Found {len(filtered_docs)} relevant documents for query")
            return filtered_docs
            
        except Exception as e:
            logger.error(f"Search error: {e}")
            return []
    
    def generate_answer(self, query: str, context_docs: List[Document]) -> Dict[str, Any]:
        """
        Tạo câu trả lời với metadata chi tiết
        """
        if not context_docs:
            return {
                'answer': "🤔 Tôi không tìm thấy thông tin liên quan trong tài liệu để trả lời câu hỏi này.",
                'confidence': 0,
                'sources_used': 0
            }
        
        # Tạo context từ documents
        context_parts = []
        for i, doc in enumerate(context_docs):
            source_name = doc.metadata.get('source', f'Tài liệu {i+1}')
            content = doc.page_content[:800]  # Limit content length
            context_parts.append(f"[Nguồn {i+1}: {source_name}]\n{content}")
        
        context = "\n\n".join(context_parts)
        
        # Tạo prompt với instructions rõ ràng
        prompt = f"""
Bạn là một trợ lý AI thông minh, chuyên phân tích và trả lời câu hỏi dựa trên tài liệu được cung cấp.

NGUYÊN TẮC TRẢ LỜI:
1. Chỉ trả lời dựa trên thông tin có trong tài liệu
2. Nếu không có đủ thông tin, hãy nói rõ
3. Trích dẫn nguồn khi cần thiết
4. Trả lời bằng tiếng Việt, rõ ràng và chi tiết
5. Nếu có nhiều quan điểm, hãy tổng hợp

TÀI LIỆU THAM KHẢO:
{context}

CÂU HỎI: {query}

TRẢ LỜI CHI TIẾT:
"""
        
        try:
            response = self.model.generate_content(prompt)
            answer = response.text
            
            # Tính confidence score đơn giản
            confidence = min(len(context_docs) * 0.2, 1.0)
            
            return {
                'answer': answer,
                'confidence': confidence,
                'sources_used': len(context_docs)
            }
            
        except Exception as e:
            logger.error(f"Answer generation error: {e}")
            return {
                'answer': f"❌ Lỗi khi tạo câu trả lời: {str(e)}",
                'confidence': 0,
                'sources_used': 0
            }
    
    def chat(self, query: str) -> Dict[str, Any]:
        """
        Xử lý câu hỏi chat với analytics
        """
        start_time = time.time()
        
        # Tìm kiếm documents liên quan
        similar_docs = self.search_similar_documents(query)
        
        # Tạo câu trả lời
        response_data = self.generate_answer(query, similar_docs)
        
        # Thông tin sources
        sources = []
        for doc in similar_docs:
            sources.append({
                'file_name': doc.metadata.get('source', 'Không rõ'),
                'file_id': doc.metadata.get('file_id', ''),
                'web_view_link': doc.metadata.get('web_view_link', ''),
                'content_preview': doc.page_content[:200] + "..." if len(doc.page_content) > 200 else doc.page_content,
                'content_length': doc.metadata.get('content_length', 0)
            })
        
        # Analytics
        processing_time = time.time() - start_time
        
        return {
            'answer': response_data['answer'],
            'sources': sources,
            'num_sources': len(sources),
            'confidence': response_data['confidence'],
            'processing_time': processing_time,
            'query_length': len(query)
        }

# Streamlit UI với cải tiến
def main():
    st.set_page_config(
        page_title="RAG Chatbot - Google Drive & Gemini",
        page_icon="🤖",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # Custom CSS
    st.markdown("""
    <style>
    .main-header {
        text-align: center;
        padding: 1rem 0;
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        color: white;
        border-radius: 10px;
        margin-bottom: 2rem;
    }
    .metric-card {
        background: #f0f2f6;
        padding: 1rem;
        border-radius: 10px;
        margin: 0.5rem 0;
    }
    .source-card {
        background: #ffffff;
        padding: 1rem;
        border-left: 4px solid #667eea;
        margin: 0.5rem 0;
        border-radius: 5px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Header
    st.markdown("""
    <div class="main-header">
        <h1>🤖 RAG Chatbot</h1>
        <p>Trợ lý AI thông minh với Google Drive & Gemini</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Sidebar cấu hình
    with st.sidebar:
        st.header("⚙️ Cấu hình")
        
        # API Key
        gemini_api_key = st.text_input(
            "🔑 Gemini API Key",
            type="password",
            value=os.getenv('GEMINI_API_KEY', ''),
            help="Nhập API key của Gemini hoặc đặt trong file .env"
        )
        
        # Folder ID
        folder_id = st.text_input(
            "📁 Google Drive Folder ID (tùy chọn)",
            value=os.getenv('GOOGLE_DRIVE_FOLDER_ID', ''),
            help="ID của folder cần đọc file. Để trống để đọc tất cả file"
        )
        
        # Advanced settings
        with st.expander("🔧 Cài đặt nâng cao"):
            max_files = st.slider("Số file tối đa", 10, 200, 50)
            chunk_size = st.slider("Kích thước chunk", 500, 2000, 1000)
            max_search_results = st.slider("Số kết quả tìm kiếm", 3, 10, 5)
        
        # Loại file
        file_types = st.multiselect(
            "📄 Loại file cần đọc",
            ['pdf', 'docx', 'txt', 'xlsx'],
            default=['pdf', 'docx', 'txt'],
            help="Chọn các loại file muốn xử lý"
        )
        
        st.markdown("---")
        
        # Stats nếu đã load
        if 'chatbot' in st.session_state and st.session_state.get('documents_loaded'):
            st.markdown("### 📊 Thống kê")
            if 'file_stats' in st.session_state:
                stats = st.session_state.file_stats
                st.metric("📁 Files đã tải", stats.get('total_files', 0))
                st.metric("📝 Đoạn text", stats.get('total_chunks', 0))
                st.metric("💾 Cache hits", len(st.session_state.chatbot.file_cache))
        
        st.markdown("---")
        
        # Quick actions
        if st.button("🗑️ Xóa cache"):
            if os.path.exists(CACHE_DIR):
                import shutil
                shutil.rmtree(CACHE_DIR)
                os.makedirs(CACHE_DIR, exist_ok=True)
                st.success("Cache đã được xóa")
                st.experimental_rerun()
        
        if st.button("🔄 Reset ứng dụng"):
            for key in list(st.session_state.keys()):
                del st.session_state[key]
            st.experimental_rerun()
    
    # Initialize session state
    if 'chatbot' not in st.session_state:
        st.session_state.chatbot = None
    if 'messages' not in st.session_state:
        st.session_state.messages = []
    if 'documents_loaded' not in st.session_state:
        st.session_state.documents_loaded = False
    
    # Main content area
    col1, col2 = st.columns([2, 1])
    
    with col1:
        # Kết nối và tải tài liệu
        if st.button("🚀 Kết nối & Tải tài liệu", type="primary"):
            if not gemini_api_key:
                st.error("❌ Vui lòng nhập Gemini API Key")
            else:
                with st.spinner("🔄 Đang khởi tạo hệ thống..."):
                    try:
                        # Khởi tạo chatbot
                        chatbot = GoogleDriveRAGChatbot(gemini_api_key)
                        
                        # Cập nhật cấu hình từ UI
                        chatbot.chunk_size = chunk_size
                        chatbot.max_search_results = max_search_results
                        
                        # Kết nối Google Drive
                        if chatbot.authenticate_google_drive():
                            st.success("✅ Kết nối Google Drive thành công!")
                            
                            # Lấy danh sách files
                            files = chatbot.get_files_from_drive(
                                folder_id if folder_id else None, 
                                file_types, 
                                max_files
                            )
                            
                            if files:
                                st.info(f"📁 Tìm thấy {len(files)} file(s) phù hợp")
                                
                                # Hiển thị preview files
                                with st.expander("👀 Preview files sẽ được xử lý"):
                                    for file_info in files[:10]:  # Hiển thị 10 đầu
                                        size_mb = int(file_info.get('size', 0)) / (1024*1024)
                                        st.write(f"📄 **{file_info['name']}** ({size_mb:.1f}MB)")
                                    if len(files) > 10:
                                        st.write(f"... và {len(files)-10} file khác")
                                
                                # Xử lý documents
                                documents = chatbot.process_documents(files)
                                
                                if documents:
                                    # Tạo vector store
                                    chatbot.create_vector_store(documents)
                                    
                                    # Lưu vào session state
                                    st.session_state.chatbot = chatbot
                                    st.session_state.documents_loaded = True
                                    
                                    # Lưu stats
                                    st.session_state.file_stats = {
                                        'total_files': len(files),
                                        'total_documents': len(documents),
                                        'total_chunks': sum(len(chatbot.text_splitter.split_text(doc.page_content)) for doc in documents),
                                        'total_size_mb': sum(int(f.get('size', 0)) for f in files) / (1024*1024)
                                    }
                                    
                                    st.success("🎉 Hệ thống đã sẵn sàng để chat!")
                                    
                                    # Hiển thị thống kê
                                    stats = st.session_state.file_stats
                                    col_stat1, col_stat2, col_stat3 = st.columns(3)
                                    with col_stat1:
                                        st.metric(" Files", stats['total_files'])
                                    with col_stat2:
                                        st.metric(" Documents", stats['total_documents'])
                                    with col_stat3:
                                        st.metric(" Chunks", stats['total_chunks'])
                                    
                                else:
                                    st.error("❌ Không thể xử lý tài liệu")
                            else:
                                st.warning("⚠️ Không tìm thấy file nào phù hợp")
                        else:
                            st.error("❌ Không thể kết nối Google Drive")
                            
                    except Exception as e:
                        st.error(f"❌ Lỗi: {str(e)}")
                        logger.error(f"Setup error: {e}")
    
    with col2:
        # Trạng thái hệ thống
        st.markdown("### 🎯 Trạng thái hệ thống")
        
        if st.session_state.documents_loaded:
            st.success("🟢 Đã sẵn sàng")
            if 'file_stats' in st.session_state:
                stats = st.session_state.file_stats
                st.markdown(f"""
                <div class="metric-card">
                    <strong>📊 Thống kê:</strong><br>
                    📁 Files: {stats.get('total_files', 0)}<br>
                    📄 Documents: {stats.get('total_documents', 0)}<br>
                    📝 Chunks: {stats.get('total_chunks', 0)}<br>
                    💾 Size: {stats.get('total_size_mb', 0):.1f} MB
                </div>
                """, unsafe_allow_html=True)
        else:
            st.warning("🟡 Chưa kết nối")
            st.info("👆 Nhấn nút 'Kết nối & Tải tài liệu' để bắt đầu")
        
        # Hướng dẫn nhanh
        with st.expander("💡 Mẹo sử dụng"):
            st.markdown("""
            **Câu hỏi hiệu quả:**
            - Hỏi cụ thể về nội dung tài liệu
            - Yêu cầu tóm tắt hoặc phân tích
            - Tìm kiếm thông tin chi tiết
            
            **Ví dụ:**
            - "Tóm tắt nội dung chính"
            - "Tìm thông tin về..."
            - "So sánh giữa... và..."
            """)
    
    # Chat interface
    if st.session_state.documents_loaded and st.session_state.chatbot:
        st.markdown("---")
        st.header("💬 Trò chuyện với AI")
        
        # Chat container
        chat_container = st.container()
        
        with chat_container:
            # Hiển thị lịch sử chat
            for i, message in enumerate(st.session_state.messages):
                with st.chat_message(message["role"]):
                    st.write(message["content"])
                    
                    # Hiển thị metadata cho assistant
                    if message["role"] == "assistant" and "metadata" in message:
                        metadata = message["metadata"]
                        
                        # Confidence và timing
                        col_meta1, col_meta2, col_meta3 = st.columns(3)
                        with col_meta1:
                            confidence = metadata.get('confidence', 0)
                            st.caption(f"🎯 Độ tin cậy: {confidence:.1%}")
                        with col_meta2:
                            processing_time = metadata.get('processing_time', 0)
                            st.caption(f"⏱️ Thời gian: {processing_time:.1f}s")
                        with col_meta3:
                            num_sources = metadata.get('num_sources', 0)
                            st.caption(f"📚 Nguồn: {num_sources}")
                        
                        # Hiển thị sources
                        if metadata.get('sources'):
                            with st.expander(f"📚 Nguồn tham khảo ({len(metadata['sources'])} tài liệu)", expanded=False):
                                for j, source in enumerate(metadata['sources'], 1):
                                    st.markdown(f"""
                                    <div class="source-card">
                                        <strong>{j}. {source['file_name']}</strong>
                                        {f'<br><a href="{source["web_view_link"]}" target="_blank">🔗 Xem file</a>' if source.get('web_view_link') else ''}
                                        <br><em>{source['content_preview']}</em>
                                    </div>
                                    """, unsafe_allow_html=True)
        
        # Input chat với suggestions
        st.markdown("### 💭 Đặt câu hỏi")
        
        # Quick suggestions
        if not st.session_state.messages:
            st.markdown("**💡 Gợi ý câu hỏi:**")
            suggestion_cols = st.columns(3)
            
            suggestions = [
                "Tóm tắt nội dung chính",
                "Những điểm quan trọng nhất",
                "Có thông tin gì về..."
            ]
            
            for i, suggestion in enumerate(suggestions):
                with suggestion_cols[i]:
                    if st.button(suggestion, key=f"suggestion_{i}"):
                        # Tự động điền suggestion vào chat
                        st.session_state.auto_query = suggestion
        
        # Chat input
        query = st.chat_input(
            "Nhập câu hỏi của bạn...",
            key="chat_input"
        )
        
        # Xử lý auto query từ suggestions
        if 'auto_query' in st.session_state:
            query = st.session_state.auto_query
            del st.session_state.auto_query
        
        if query:
            # Validate query
            if len(query.strip()) < 3:
                st.warning(" Câu hỏi quá ngắn. Vui lòng nhập câu hỏi chi tiết hơn.")
                st.stop()
            
            # Hiển thị câu hỏi
            with st.chat_message("user"):
                st.write(query)
            
            # Thêm vào lịch sử
            st.session_state.messages.append({"role": "user", "content": query})
            
            # Tạo câu trả lời
            with st.chat_message("assistant"):
                with st.spinner("🤔 Đang suy nghĩ..."):
                    try:
                        response = st.session_state.chatbot.chat(query)
                        
                        # Hiển thị câu trả lời
                        st.write(response['answer'])
                        
                        # Hiển thị metadata
                        col_meta1, col_meta2, col_meta3 = st.columns(3)
                        with col_meta1:
                            st.caption(f"Độ tin cậy: {response['confidence']:.1%}")
                        with col_meta2:
                            st.caption(f"Thời gian: {response['processing_time']:.1f}s")
                        with col_meta3:
                            st.caption(f"📚 Nguồn: {response['num_sources']}")
                        
                        # Hiển thị sources
                        if response['sources']:
                            with st.expander(f"📚 Nguồn tham khảo ({response['num_sources']} tài liệu)", expanded=False):
                                for j, source in enumerate(response['sources'], 1):
                                    st.markdown(f"""
                                    <div class="source-card">
                                        <strong>{j}. {source['file_name']}</strong>
                                        {f'<br><a href="{source["web_view_link"]}" target="_blank">🔗 Xem file</a>' if source.get('web_view_link') else ''}
                                        <br><em>{source['content_preview']}</em>
                                    </div>
                                    """, unsafe_allow_html=True)
                        
                        # Đánh giá phản hồi
                        st.markdown("---")
                        feedback_col1, feedback_col2, feedback_col3 = st.columns([1, 1, 3])
                        
                        with feedback_col1:
                            if st.button("👍", key=f"thumbs_up_{len(st.session_state.messages)}"):
                                st.success("Cảm ơn phản hồi!")
                        
                        with feedback_col2:
                            if st.button("👎", key=f"thumbs_down_{len(st.session_state.messages)}"):
                                st.info("Chúng tôi sẽ cải thiện!")
                        
                        with feedback_col3:
                            st.caption("Phản hồi giúp cải thiện chất lượng trả lời")
                        
                    except Exception as e:
                        st.error(f"❌ Lỗi khi xử lý câu hỏi: {str(e)}")
                        logger.error(f"Chat error: {e}")
                        response = {
                            'answer': "Xin lỗi, đã có lỗi xảy ra khi xử lý câu hỏi của bạn.",
                            'sources': [],
                            'num_sources': 0,
                            'confidence': 0,
                            'processing_time': 0
                        }
            
            # Thêm vào lịch sử với metadata
            st.session_state.messages.append({
                "role": "assistant",
                "content": response['answer'],
                "metadata": {
                    'sources': response['sources'],
                    'num_sources': response['num_sources'],
                    'confidence': response['confidence'],
                    'processing_time': response['processing_time']
                }
            })
    
    else:
        # Hướng dẫn khi chưa setup
        st.markdown("---")
        st.info("🚀 **Hãy kết nối với Google Drive để bắt đầu!**")
        
        # Checklist setup
        st.markdown("### ✅ Checklist chuẩn bị:")
        
        checklist_items = [
            ("📁 File credentials.json", os.path.exists(CREDENTIALS_FILE)),
            ("🔑 Gemini API Key", bool(gemini_api_key)),
            ("📄 Chọn loại file", bool(file_types)),
        ]
        
        for item, status in checklist_items:
            if status:
                st.success(f"✅ {item}")
            else:
                st.error(f"❌ {item}")
        
        # Demo queries
        if all(status for _, status in checklist_items):
            st.success("🎉 Bạn đã sẵn sàng! Nhấn nút 'Kết nối & Tải tài liệu' phía trên.")
        else:
            st.warning("⚠️ Vui lòng hoàn thành checklist trước khi tiếp tục.")

if __name__ == "__main__":
    main()
    from langchain_google_genai import GoogleGenerativeAIEmbeddings